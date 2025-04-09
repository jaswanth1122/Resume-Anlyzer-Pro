import os
import json
import streamlit as st
from dotenv import load_dotenv
import google.generativeai as genai
from pdf2image import convert_from_path
import pytesseract
import pdfplumber
from langdetect import detect
from googletrans import Translator
from fpdf import FPDF
from docx import Document

# --- Configuration ---
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
model = genai.GenerativeModel("gemini-1.5-flash")

# --- PDF Processing ---
def extract_text_from_pdf(uploaded_file):
    text = ""
    try:
        with open("temp.pdf", "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        with pdfplumber.open("temp.pdf") as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text
        
        if text.strip():
            return text.strip()
    except Exception as e:
        st.warning(f"Direct extraction failed: {e}")

    try:
        images = convert_from_path("temp.pdf")
        for image in images:
            text += pytesseract.image_to_string(image) + "\n"
        return text.strip()
    except Exception as e:
        st.error(f"OCR failed: {e}")
        return ""

# --- Language Support ---
def detect_and_translate(text, target_lang='en'):
    try:
        src_lang = detect(text)
        if src_lang != target_lang:
            translator = Translator()
            translation = translator.translate(text, src=src_lang, dest=target_lang)
            return translation.text
        return text
    except:
        return text

# --- Resume Analysis ---
def analyze_resume(resume_text, job_description=None, language='en'):
    translated_text = detect_and_translate(resume_text, 'en')
    
    prompt = [
        "Analyze this resume (originally in {}) and provide:".format(language),
        "1. Skills Summary (Technical & Soft Skills)",
        "2. Missing Skills for Target Roles",
        "3. Improvement Recommendations (Courses/Certifications)",
        "4. Strengths & Weaknesses",
        "5. {}".format("Job Fit Analysis" if job_description else "General Career Advice"),
        "",
        "Resume:",
        translated_text
    ]
    
    if job_description:
        prompt.extend([
            "",
            "Job Description:",
            job_description,
            "",
            "Compare and highlight:",
            "- Matching qualifications",
            "- Missing requirements",
            "- Suggested adaptations"
        ])
    
    response = model.generate_content("\n".join(prompt))
    return response.text

# --- ATS Compliance Check ---
def check_ats_compliance(resume_text):
    prompt = "\n".join([
        "Evaluate ATS (Applicant Tracking System) compliance:",
        "1. Keyword Optimization",
        "2. Proper Headings",
        "3. No Graphics/Tables",
        "4. Standard Fonts",
        "5. Correct File Format",
        "6. Skills Match",
        "7. Contact Info Visibility",
        "",
        "Return JSON with:",
        '{"strengths": [str],',
        '"weaknesses": [str],',
        '"suggestions": [str]}',
        "",
        "Resume: {}".format(resume_text[:2000])
    ])
    
    try:
        response = model.generate_content(prompt)
        return json.loads(response.text)
    except:
        return {"strengths": [], "weaknesses": [], "suggestions": []}

# --- Report Generation (Fixed for Windows) ---
def clean_text_for_pdf(text):
    replacements = {
        "–": "-",
        "—": "-",
        "‘": "'",
        "’": "'",
        "“": '"',
        "”": '"',
        "…": "..."
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text

def generate_pdf_report(content, filename="report.pdf"):
    try:
        pdf = FPDF()
        pdf.add_page()
        
        # First try with Arial (common on Windows)
        try:
            pdf.add_font("Arial", "", "c:/windows/fonts/arial.ttf", uni=True)
            pdf.set_font("Arial", size=12)
        except:
            # Fallback to Helvetica with cleaned text
            content = clean_text_for_pdf(content)
            pdf.set_font("Helvetica", size=12)
        
        pdf.multi_cell(0, 10, content)
        pdf.output(filename)
        return filename
    except Exception as e:
        st.error(f"PDF generation failed: {str(e)}")
        return None

def generate_word_report(content, filename="report.docx"):
    try:
        doc = Document()
        doc.add_heading('Resume Analysis Report', 0)
        for line in content.split('\n'):
            if line.strip():  # Skip empty lines
                doc.add_paragraph(line)
        doc.save(filename)
        return filename
    except Exception as e:
        st.error(f"Word report generation failed: {str(e)}")
        return None

# --- Streamlit UI ---
st.set_page_config(page_title="AI Resume Analyzer Pro", layout="wide")
st.title("📄 AI-Powered Resume Analysis")
st.write("Get comprehensive resume feedback with AI")

with st.sidebar:
    st.header("Settings")
    language = st.selectbox("Resume Language", ["en", "es", "fr", "de"])
    analysis_depth = st.radio("Analysis Depth", ["Basic", "Detailed", "Comprehensive"])

col1, col2 = st.columns(2)

with col1:
    uploaded_file = st.file_uploader("Upload Resume (PDF)", type=["pdf"])

with col2:
    job_desc = st.text_area("Paste Job Description (Optional)", height=200)

if uploaded_file:
    with st.spinner("Processing resume..."):
        resume_text = extract_text_from_pdf(uploaded_file)
        
        if st.button("Run Full Analysis"):
            with st.spinner("Generating insights..."):
                analysis = analyze_resume(resume_text, job_desc, language)
                ats_data = check_ats_compliance(resume_text)
                
                st.success("Analysis Complete!")
                
                # ATS Compliance
                st.subheader("📋 ATS Compliance Analysis")
                with st.expander("ATS Details"):
                    st.write("✅ Strengths:")
                    for item in ats_data['strengths']:
                        st.write(f"- {item}")
                    
                    st.write("⚠️ Improvements Needed:")
                    for item in ats_data['suggestions']:
                        st.write(f"- {item}")
                
                # Detailed Analysis
                st.subheader("📝 Detailed Analysis")
                st.write(analysis)
                
                # Report Generation
                report_content = "\n".join([
                    "RESUME ANALYSIS REPORT",
                    "----------------------",
                    "ATS COMPLIANCE ANALYSIS:",
                    f"Strengths: {', '.join(ats_data['strengths'])}",
                    f"Suggestions: {', '.join(ats_data['suggestions'])}",
                    "",
                    "DETAILED ANALYSIS:",
                    analysis
                ])
                
                # Export buttons
                col1, col2 = st.columns(2)
                with col1:
                    pdf_path = generate_pdf_report(report_content)
                    if pdf_path:
                        with open(pdf_path, "rb") as f:
                            st.download_button("Download PDF Report", f, "resume_analysis.pdf")
                
                with col2:
                    docx_path = generate_word_report(report_content)
                    if docx_path:
                        with open(docx_path, "rb") as f:
                            st.download_button("Download Word Report", f, "resume_analysis.docx")

# Footer
st.markdown("---")
st.markdown("""
<p style='text-align: center;'>
  Powered by Google Gemini AI • 
  Developed By
  <a href="https://www.linkedin.com/in/jaswanthvalluri11/" target="_blank">Jaswanth Valluri</a> • 
  Made with Streamlit
</p>
""", unsafe_allow_html=True)
